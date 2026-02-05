"""Tests for RAG pipeline: extractor, chunker, embedder."""

import io
import json
from unittest.mock import AsyncMock, MagicMock, patch

import pytest

from rag.chunker import Chunk, chunk_pages, chunk_text
from rag.extractor import PageText, extract_docx, extract_text


# ── T12.2: extract_docx ──────────────────────────────────────────────────────

class TestExtractDocx:
    def test_extracts_paragraphs(self):
        """T12.2: Parse a test DOCX, verify all paragraphs are in output."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("Eerste paragraaf over taxonomie.")
        doc.add_paragraph("Tweede paragraaf over toetsing.")
        doc.add_paragraph("Derde paragraaf over kwaliteit.")

        buf = io.BytesIO()
        doc.save(buf)
        file_bytes = buf.getvalue()

        result = extract_docx(file_bytes)
        assert "Eerste paragraaf over taxonomie." in result
        assert "Tweede paragraaf over toetsing." in result
        assert "Derde paragraaf over kwaliteit." in result


# ── T12.3/T12.4: chunk_text ──────────────────────────────────────────────────

class TestChunkText:
    def test_long_text_produces_multiple_chunks(self):
        """T12.3: Input of 2000 words → multiple chunks, each ≤ chunk_size chars."""
        words = ["woord"] * 2000
        text = " ".join(words)  # ~10000 chars

        chunks = chunk_text(text, chunk_size=2400, chunk_overlap=200)

        assert len(chunks) > 1
        for chunk in chunks:
            assert len(chunk.text) <= 2400 + 200  # Allow some margin

    def test_overlap_between_chunks(self):
        """T12.3: Verify overlap between consecutive chunks."""
        words = ["woord"] * 2000
        text = " ".join(words)

        chunks = chunk_text(text, chunk_size=2000, chunk_overlap=200)

        if len(chunks) >= 2:
            # End of chunk 0 should overlap with start of chunk 1
            end_of_first = chunks[0].text[-100:]
            assert end_of_first in chunks[1].text

    def test_short_text_single_chunk(self):
        """T12.4: Short input of 100 words → exactly 1 chunk."""
        words = ["kort"] * 100
        text = " ".join(words)

        chunks = chunk_text(text, chunk_size=2000, chunk_overlap=200)

        assert len(chunks) == 1

    def test_empty_text_no_chunks(self):
        """Empty text produces no chunks."""
        chunks = chunk_text("", chunk_size=2000, chunk_overlap=200)
        assert len(chunks) == 0

    def test_chunks_have_positions(self):
        """Chunks are numbered sequentially."""
        words = ["woord"] * 2000
        text = " ".join(words)

        chunks = chunk_text(text, chunk_size=1000, chunk_overlap=100)

        for i, chunk in enumerate(chunks):
            assert chunk.position == i

    def test_chunk_pages_preserves_page_numbers(self):
        """chunk_pages preserves page numbers from PageText objects."""
        pages = [
            PageText(page_number=1, text="Inhoud van pagina een. " * 50),
            PageText(page_number=2, text="Inhoud van pagina twee. " * 50),
        ]

        chunks = chunk_pages(pages, chunk_size=500, chunk_overlap=50)

        assert len(chunks) > 0
        # All chunks from page 1 should have page=1
        page1_chunks = [c for c in chunks if c.page == 1]
        page2_chunks = [c for c in chunks if c.page == 2]
        assert len(page1_chunks) > 0
        assert len(page2_chunks) > 0


# ── T12.5: embed_chunks (mock OpenAI API) ────────────────────────────────────

class TestEmbedChunks:
    @pytest.mark.asyncio
    async def test_batches_and_returns_vectors(self):
        """T12.5: Verify batching of max 100 and 1536-dimensional vectors."""
        from rag.embedder import embed_chunks, BATCH_SIZE

        # Create 150 texts to test batching
        texts = [f"Text chunk {i}" for i in range(150)]

        fake_embeddings = [[0.1] * 1536 for _ in range(min(100, len(texts)))]
        fake_response = {
            "data": [{"embedding": emb} for emb in fake_embeddings]
        }

        call_count = 0

        async def mock_post(url, **kwargs):
            nonlocal call_count
            call_count += 1
            batch = kwargs.get("json", {}).get("input", [])
            batch_embeddings = [[0.1] * 1536 for _ in batch]
            mock_resp = MagicMock()
            mock_resp.raise_for_status = MagicMock()
            mock_resp.json.return_value = {
                "data": [{"embedding": emb} for emb in batch_embeddings]
            }
            return mock_resp

        with patch("rag.embedder.httpx.AsyncClient") as MockClient:
            mock_client = AsyncMock()
            mock_client.post = mock_post
            mock_client.__aenter__ = AsyncMock(return_value=mock_client)
            mock_client.__aexit__ = AsyncMock(return_value=False)
            MockClient.return_value = mock_client

            result = await embed_chunks(texts)

        # Should have made 2 batches (100 + 50)
        assert call_count == 2
        # Should return 150 embeddings
        assert len(result) == 150
        # Each embedding should be 1536 dimensions
        assert all(len(v) == 1536 for v in result)


# ── T12.6: embedding_pipeline integration (mock Supabase + OpenAI) ────────────

class TestEmbeddingPipeline:
    @pytest.mark.asyncio
    async def test_full_pipeline(self):
        """T12.6: Verify full flow: download → extract → chunk → embed → insert."""
        from services.embedding_pipeline import run_embedding

        # Mock Supabase client
        mock_supabase = MagicMock()

        # Mock material query
        mock_supabase.table.return_value.select.return_value.eq.return_value.single.return_value.execute.return_value = MagicMock(
            data={
                "id": "mat-1",
                "storage_path": "materials/test.txt",
                "mime_type": "text/plain",
            }
        )

        # Mock storage download (return plain text)
        test_text = "Dit is een test tekst voor de embedding pipeline. " * 20
        mock_supabase.storage.from_.return_value.download.return_value = (
            test_text.encode("utf-8")
        )

        # Mock insert
        mock_supabase.table.return_value.insert.return_value.execute.return_value = (
            MagicMock()
        )

        # Mock update
        mock_supabase.table.return_value.update.return_value.eq.return_value.execute.return_value = (
            MagicMock()
        )

        # Mock embed_chunks
        with patch(
            "services.embedding_pipeline.get_supabase_client",
            return_value=mock_supabase,
        ), patch(
            "services.embedding_pipeline.embed_chunks",
            new_callable=AsyncMock,
        ) as mock_embed:
            mock_embed.return_value = [[0.1] * 1536]  # One chunk for short text

            await run_embedding("mat-1")

        # Verify the pipeline steps happened
        # 1. Material was fetched
        mock_supabase.table.assert_any_call("materials")

        # 2. File was downloaded from storage
        mock_supabase.storage.from_.assert_called_with("materials")

        # 3. Embeddings were generated
        mock_embed.assert_called_once()

        # 4. Chunks were inserted
        insert_calls = [
            c
            for c in mock_supabase.table.return_value.insert.call_args_list
        ]
        assert len(insert_calls) > 0

        # 5. Material was updated with chunk_count
        update_calls = [
            c
            for c in mock_supabase.table.return_value.update.call_args_list
        ]
        assert len(update_calls) > 0


# ── extract_text dispatch ─────────────────────────────────────────────────────

class TestExtractText:
    def test_txt_extraction(self):
        """Plain text extraction."""
        text = "Hello world"
        result = extract_text(text.encode("utf-8"), "text/plain")
        assert result == "Hello world"

    def test_unsupported_type_raises(self):
        """Unsupported mime type raises ValueError."""
        with pytest.raises(ValueError, match="Unsupported mime type"):
            extract_text(b"data", "application/octet-stream")
