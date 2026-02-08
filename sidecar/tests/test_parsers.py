"""T8.1-T8.4, T8.7: Parser tests for CSV, XLSX, DOCX, and /parse endpoint."""

import io
from pathlib import Path

import pytest
from fastapi.testclient import TestClient

from main import app
from parsers.csv_parser import parse_csv
from parsers.xlsx_parser import parse_xlsx
from parsers.docx_parser import parse_docx

FIXTURES = Path(__file__).parent / "fixtures"

client = TestClient(app)


class TestCsvParser:
    """T8.1 & T8.2: CSV parser tests."""

    def test_parse_5_questions(self):
        """T8.1: Parse a test CSV with 5 questions."""
        content = (FIXTURES / "test_questions.csv").read_text()
        questions = parse_csv(content)
        assert len(questions) == 5

        # First question
        assert questions[0].stem == "Wat is de hoofdstad van Nederland?"
        assert len(questions[0].options) == 4
        assert questions[0].options[0].is_correct  # A = Amsterdam
        assert questions[0].options[0].text == "Amsterdam"
        assert not questions[0].options[1].is_correct

        # Third question (correct = C)
        assert questions[2].stem == "Wat is H2O?"
        assert questions[2].options[2].is_correct  # C = Water
        assert questions[2].options[2].text == "Water"

    def test_semicolon_delimiter(self):
        content = (FIXTURES / "test_questions_semicolon.csv").read_text()
        questions = parse_csv(content)
        assert len(questions) == 1
        assert questions[0].stem == "Wat is 2+2?"
        assert questions[0].options[1].is_correct  # B = 4

    def test_missing_column_raises_error(self):
        """T8.2: CSV with missing column raises clear error."""
        content = (FIXTURES / "test_missing_column.csv").read_text()
        with pytest.raises(ValueError, match="Ontbrekende kolommen"):
            parse_csv(content)

    def test_empty_stem_row_is_included(self):
        """Rows with missing stem should be parsed (not skipped) so validation can flag them."""
        content = (
            "stam,optie_a,optie_b,optie_c,optie_d,correct\n"
            "Vraag 1?,A,B,C,D,A\n"
            ",Opt A,Opt B,Opt C,Opt D,B\n"
            "Vraag 3?,X,Y,Z,W,C\n"
        )
        questions = parse_csv(content)
        assert len(questions) == 3
        assert questions[0].stem == "Vraag 1?"
        assert questions[1].stem == ""
        assert questions[2].stem == "Vraag 3?"

    def test_invalid_correct_is_not_fatal(self):
        """Invalid 'correct' value should not crash the parser; validation catches it."""
        content = (
            "stam,optie_a,optie_b,optie_c,optie_d,correct\n"
            "Vraag 1?,A,B,C,D,A\n"
            "Vraag 2?,X,Y,Z,W,E\n"
        )
        questions = parse_csv(content)
        assert len(questions) == 2
        # Second question: no option marked correct
        assert not any(o.is_correct for o in questions[1].options)

    def test_completely_empty_row_is_skipped(self):
        """A row where all fields are blank should still be skipped."""
        content = (
            "stam,optie_a,optie_b,optie_c,optie_d,correct\n"
            "Vraag 1?,A,B,C,D,A\n"
            ",,,,, \n"
            "Vraag 3?,X,Y,Z,W,C\n"
        )
        questions = parse_csv(content)
        assert len(questions) == 2

    def test_bytes_input(self):
        content = (FIXTURES / "test_questions.csv").read_bytes()
        questions = parse_csv(content)
        assert len(questions) == 5


class TestXlsxParser:
    """T8.3: XLSX parser tests."""

    @pytest.fixture
    def xlsx_bytes(self) -> bytes:
        """Create a test XLSX file in memory."""
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.append(["stam", "optie_a", "optie_b", "optie_c", "optie_d", "correct"])
        ws.append(["Wat is 1+1?", "1", "2", "3", "4", "B"])
        ws.append(["Wat is 2+2?", "3", "4", "5", "6", "B"])
        ws.append(["Wat is 3+3?", "5", "6", "7", "8", "B"])

        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue()

    def test_parse_3_questions(self, xlsx_bytes: bytes):
        """T8.3: Parse a test Excel with 3 questions."""
        questions = parse_xlsx(xlsx_bytes)
        assert len(questions) == 3
        assert questions[0].stem == "Wat is 1+1?"
        assert questions[0].options[1].is_correct  # B = 2
        assert questions[0].options[1].text == "2"

    def test_empty_stem_row_is_included(self):
        """Rows with missing stem should be parsed (not skipped) so validation can flag them."""
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.append(["stam", "optie_a", "optie_b", "optie_c", "optie_d", "correct"])
        ws.append(["Vraag 1?", "A", "B", "C", "D", "A"])
        ws.append([None, "Opt A", "Opt B", "Opt C", "Opt D", "B"])
        ws.append(["Vraag 3?", "X", "Y", "Z", "W", "C"])

        buf = io.BytesIO()
        wb.save(buf)
        questions = parse_xlsx(buf.getvalue())
        assert len(questions) == 3
        assert questions[1].stem == ""

    def test_invalid_correct_is_not_fatal(self):
        """Invalid 'correct' value should not crash the XLSX parser."""
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.append(["stam", "optie_a", "optie_b", "optie_c", "optie_d", "correct"])
        ws.append(["Vraag?", "A", "B", "C", "D", "E"])

        buf = io.BytesIO()
        wb.save(buf)
        questions = parse_xlsx(buf.getvalue())
        assert len(questions) == 1
        assert not any(o.is_correct for o in questions[0].options)


class TestDocxParser:
    """T8.4: DOCX parser tests."""

    @pytest.fixture
    def docx_bytes(self) -> bytes:
        """Create a test DOCX with bold correct answers."""
        from docx import Document
        from docx.shared import Pt

        doc = Document()
        doc.add_paragraph("1. Wat is de hoofdstad van Nederland?")
        doc.add_paragraph("A. Rotterdam")
        # Bold for correct answer
        p = doc.add_paragraph()
        run = p.add_run("B. Amsterdam")
        run.bold = True
        doc.add_paragraph("C. Utrecht")
        doc.add_paragraph("D. Den Haag")

        doc.add_paragraph("")
        doc.add_paragraph("2. Welk dier is een zoogdier?")
        doc.add_paragraph("A. Haai")
        doc.add_paragraph("B. Slang")
        doc.add_paragraph("C. Hond *")
        doc.add_paragraph("D. Krokodil")

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def test_parse_docx_with_bold_correct(self, docx_bytes: bytes):
        """T8.4: Parse DOCX with numbered questions and bold correct answer."""
        questions = parse_docx(docx_bytes)
        assert len(questions) == 2

        # Question 1: B is bold (correct)
        assert questions[0].stem == "Wat is de hoofdstad van Nederland?"
        assert len(questions[0].options) == 4
        assert questions[0].options[1].is_correct  # B = Amsterdam

        # Question 2: C has asterisk (correct)
        assert questions[1].stem == "Welk dier is een zoogdier?"
        assert questions[1].options[2].is_correct  # C = Hond


class TestParseEndpoint:
    """T8.7: Sidecar /parse integration test."""

    def test_parse_csv_endpoint(self):
        csv_content = (FIXTURES / "test_questions.csv").read_bytes()
        response = client.post(
            "/parse",
            files={"file": ("test.csv", csv_content, "text/csv")},
        )
        assert response.status_code == 200
        data = response.json()
        assert len(data) == 5
        assert data[0]["stem"] == "Wat is de hoofdstad van Nederland?"
        assert len(data[0]["options"]) == 4

    def test_parse_unsupported_format(self):
        response = client.post(
            "/parse",
            files={"file": ("test.txt", b"hello", "text/plain")},
        )
        assert response.status_code == 400

    def test_parse_invalid_csv(self):
        response = client.post(
            "/parse",
            files={
                "file": (
                    "bad.csv",
                    b"stam,optie_a,correct\ntest,a,A",
                    "text/csv",
                )
            },
        )
        assert response.status_code == 422
